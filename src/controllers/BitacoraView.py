# /src/views/GiroView

from flask import Flask, request, json, Response, Blueprint, g
from marshmallow import ValidationError
from sqlalchemy import true
from flask import send_file
from ..models.ProyectoModel import ProyectoModel,ProyectoSchema,ProyectoSchemaUpdate,ProyectosSchemaQuery
from ..models.BitacoraModel import BitacoraModel,BitacoraSchema,BitacoraSchemaQuery,BitacoraSchemaUpdate
from ..models.UsersModel import UsersModel
from ..models.EventoModel import EventoModel
from ..models import db
from ..shared import returnCodes
from flask_restx import Api,fields,Resource,reqparse
from io import BytesIO
from io import StringIO
from docx import Document
import csv
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Inches



app = Flask(__name__)
parser = reqparse.RequestParser()
parser.add_argument('limit', type=int, location='args')
parser.add_argument('offset', type=int, location='args')
bitacora_api = Blueprint("bitacora_api", __name__)
bitacora_schema = BitacoraSchema()
bitacora_schema_update = BitacoraSchemaUpdate()
bitacora_schema_query = BitacoraSchemaQuery()
api = Api(bitacora_api)

nsBitacora = api.namespace("bitacora", description="API operations for bitacora")

BitacoraQueryModel = nsBitacora.model(
    "bitacoraQuery",
    {
     
        "id": fields.Integer(description="identificador"),
        "proyectId": fields.Integer(description="proyectId"),
        "usuarioId" : fields.Integer(description="usuarioId"),
        "comentario" : fields.String(description="comentario"),
        "fechaInicio" : fields.DateTime(description="fechaInicio"),
        "fechaFin" : fields.DateTime(description="fechaFin"),
        "IDEvento" : fields.String(description="IDEvento"),
        "isVentana" : fields.Boolean(description="isVentana")

    }
)

#parser.add_argument('DevicesQueryModel', type=json, location='body')

BitacoraModelApi = nsBitacora.model(
    "bitacora",
    {

        "proyectId": fields.Integer(description="proyectId"),
        "usuarioId" : fields.Integer(description="usuarioId"),
        "comentario" : fields.String(description="comentario"),
        "fechaInicio" : fields.DateTime(description="fechaInicio"),
        "fechaFin" : fields.DateTime(description="fechaFin"),
        "IDEvento" : fields.String(description="IDEvento"),
        "isVentana" : fields.Boolean(description="isVentana")

    }
)


BitacoraPatchApi = nsBitacora.model(
    "bitacoraPatch",
    {
        "id": fields.Integer(description="identificador"),
        "proyectId": fields.Integer(description="proyectId"),
        "usuarioId" : fields.Integer(description="usuarioId"),
        "comentario" : fields.String(description="comentario"),
        "fechaInicio" : fields.DateTime(description="fechaInicio"),
        "fechaFin" : fields.DateTime(description="fechaFin"),
        "IDEvento" : fields.String(description="IDEvento"),
        "isVentana" : fields.Boolean(description="isVentana")
        
    }
)

def createBitacora(req_data, listaObjetosCreados, listaErrores):
    #app.logger.info("Creando catalogo" + json.dumps(req_data))



    proyecto_in_db = ProyectoModel.get_one_project(req_data.get("proyectId"))
    if not proyecto_in_db:
        #error = returnCodes.custom_response(None, 409, "TPM-5", "", data.get("nombre")).json
        error = returnCodes.partial_response("TPM-4","",req_data.get("proyectId"))
        listaErrores.append(error)
        return returnCodes.custom_response(None, 409, "TPM-4", "", req_data.get("proyectId"))

    user_in_db = UsersModel.get_one_users(req_data.get("usuarioId"))
    if not user_in_db:
        #error = returnCodes.custom_response(None, 409, "TPM-5", "", data.get("nombre")).json
        error = returnCodes.partial_response("TPM-4","",req_data.get("usuarioId"))
        listaErrores.append(error)
        return returnCodes.custom_response(None, 409, "TPM-4", "", req_data.get("usuarioId"))
    
    #check evento
    evento_in_db = EventoModel.get_evento_by_nombre_and_proyect(req_data.get("IDEvento"),req_data.get("proyectId"))
    if not evento_in_db:
        error = returnCodes.partial_response("TPM-4","",req_data.get("IDEvento"))
        listaErrores.append(error)
        return returnCodes.custom_response(None, 409, "TPM-4", "", req_data.get("IDEvento"))
    if evento_in_db.activo is False:
        error = returnCodes.partial_response("TPM-20","",req_data.get("IDEvento"))
        listaErrores.append(error)
        return returnCodes.custom_response(None, 409, "TPM-20", "", req_data.get("IDEvento"))



    bitacora = BitacoraModel(req_data)

    try:
        bitacora.save()
        
    except Exception as err:
        error = returnCodes.custom_response(None, 500, "TPM-7", str(err)).json
        error = returnCodes.partial_response("TPM-7",str(err))
        #error="Error al intentar dar de alta el registro "+data.get("nombre")+", "+error["message"]
        listaErrores.append(error)
        db.session.rollback()
        return returnCodes.custom_response(None, 500, "TPM-7", str(err))
    
    serialized_bitacora = bitacora_schema.dump(bitacora)
    listaObjetosCreados.append(serialized_bitacora)
    return returnCodes.custom_response(serialized_bitacora, 201, "TPM-1")

@nsBitacora.route("")
class BitacoraList(Resource):
    @nsBitacora.doc("lista de proyectos")
    @nsBitacora.expect(parser)
    def get(self):
        """List all status"""
        offset = 0
        limit = 10
        if "offset" in request.args:
            offset = request.args.get('offset',default = 0, type = int)

        if "limit" in request.args:
            limit = request.args.get('limit',default = 10, type = int)
        proyectos = BitacoraModel.get_all_Bitacora(offset,limit)
        #return catalogos
        serialized_bitacora = bitacora_schema.dump(proyectos, many=True)
        return returnCodes.custom_response(serialized_bitacora, 200, "TPM-3")

    @nsBitacora.doc("Crear proyecto")
    @nsBitacora.expect(BitacoraModelApi)
    @nsBitacora.response(201, "created")
    def post(self):
        if request.is_json is False:
            return returnCodes.custom_response(None, 400, "TPM-2")

        req_data = request.json
        if(not req_data):
            return returnCodes.custom_response(None, 400, "TPM-2")
        try:
            data = bitacora_schema.load(req_data)
        except ValidationError as err:    
            return returnCodes.custom_response(None, 400, "TPM-2", str(err))
        
        listaObjetosCreados = list()
        listaErrores = list()
        
       
        createBitacora(data, listaObjetosCreados, listaErrores)
        
        if(len(listaObjetosCreados)>0):
            if(len(listaErrores)==0):
                return returnCodes.custom_response(listaObjetosCreados[0], 201, "TPM-8")
            else:
                return returnCodes.custom_response(listaObjetosCreados[0], 201, "TPM-16", "",listaErrores)
        else:
            return returnCodes.custom_response(None, 409, "TPM-16","", listaErrores)
    
    @nsBitacora.doc("actualizar bitacora")
    @nsBitacora.expect(BitacoraPatchApi)
    def put(self):
        if request.is_json is False:
            return returnCodes.custom_response(None, 400, "TPM-2")
        req_data = request.json
        data = None
        if(not req_data):
            return returnCodes.custom_response(None, 400, "TPM-2")
        try:
            data = bitacora_schema_update.load(req_data, partial=True)
        except ValidationError as err:
            return returnCodes.custom_response(None, 400, "TPM-2", str(err))

        bitacora = BitacoraModel.get_one_Bitacora(data.get("id"))
        if not bitacora:
            return returnCodes.custom_response(None, 404, "TPM-4","el proyecto no existe")
        if "proyectId" in data:
            device_in_db = ProyectoModel.get_one_project(data.get("proyectId"))
            if not device_in_db:
                #error = returnCodes.custom_response(None, 409, "TPM-5", "", data.get("nombre")).json
                
                return returnCodes.custom_response(None, 409, "TPM-4", "", data.get("proyectId"))

        if "usuarioId" in data:
            user_in_db = UsersModel.get_one_users(data.get("usuarioId"))
            if not user_in_db:
                #error = returnCodes.custom_response(None, 409, "TPM-5", "", data.get("nombre")).json
                
                return returnCodes.custom_response(None, 409, "TPM-4", "", data.get("usuarioId"))

        try:
            bitacora.update(data)
        except Exception as err:
            return returnCodes.custom_response(None, 500, "TPM-7", str(err))

        serialized_bitacora = bitacora_schema.dump(bitacora)
        return returnCodes.custom_response(serialized_bitacora, 200, "TPM-6")

@nsBitacora.route("/<int:id>")
@nsBitacora.param("id", "The id identifier")
@nsBitacora.response(404, "bitacora no encontrada")
class Oneproyecto(Resource):
    @nsBitacora.doc("obtener un equipo")
    def get(self, id):
       
        proyecto = BitacoraModel.get_one_Bitacora(id)
        if not proyecto:
            return returnCodes.custom_response(None, 404, "TPM-4")

        serialized_bitacora = bitacora_schema.dump(proyecto)
        return returnCodes.custom_response(serialized_bitacora, 200, "TPM-3")

@nsBitacora.route("/query")
@nsBitacora.expect(parser)
@nsBitacora.response(404, "bitacora no encontrada")
class bitacoraQuery(Resource):
    
    @nsBitacora.doc("obtener varias bitacoras")
    @api.expect(BitacoraQueryModel)
    def post(self):
        print(request.args)
        offset = 0
        limit = 10

        if request.is_json is False:
            return returnCodes.custom_response(None, 400, "TPM-2")

        if "offset" in request.args:
            offset = request.args.get('offset',default = 0, type = int)

        if "limit" in request.args:
            limit = request.args.get('limit',default = 10, type = int)

        req_data = request.json
        data = None
        if(not req_data):
            return returnCodes.custom_response(None, 400, "TPM-2")
        try:
            data = bitacora_schema_query.load(req_data, partial=True)
        except ValidationError as err:
            return returnCodes.custom_response(None, 400, "TPM-2", str(err))

        devices = BitacoraModel.get_Bitacora_by_query(data,offset,limit)
        if not devices:
            return returnCodes.custom_response(None, 404, "TPM-4")

        serialized_bitacora = bitacora_schema.dump(devices.items,many=true)
        return returnCodes.custom_response(serialized_bitacora, 200, "TPM-3")


@nsBitacora.route("/filter/<value>")
@nsBitacora.expect(parser)
@nsBitacora.response(404, "bitacora no encontrada")
class BitacoraFilter(Resource):
    
    @nsBitacora.doc("obtener varioas bitacoras")
    def get(self,value):
      
        offset = 1
        limit = 100

        

        if "offset" in request.args:
            offset = request.args.get('offset',default = 1, type = int)

        if "limit" in request.args:
            limit = request.args.get('limit',default = 100, type = int)


        bitacoras = BitacoraModel.get_Bitacora_by_query(value,offset,limit)
        if not bitacoras:
            return returnCodes.custom_response(None, 404, "TPM-4")

        serialized_bitacora = bitacora_schema.dump(bitacoras.items,many=True)
        return returnCodes.custom_response(serialized_bitacora, 200, "TPM-3")
    

@nsBitacora.route("/csv")
@nsBitacora.expect(parser)
@nsBitacora.response(404, "bitacora no encontrada")
class bitacoraCSV(Resource):
    
    @nsBitacora.doc("obtener varias bitacoras")
    @api.expect(BitacoraQueryModel)
    def post(self):
        
        if request.is_json is False:
            return returnCodes.custom_response(None, 400, "TPM-2")

        req_data = request.json
        data = None
        if(not req_data):
            return returnCodes.custom_response(None, 400, "TPM-2")
        try:
            data = bitacora_schema_query.load(req_data, partial=True)
        except ValidationError as err:
            return returnCodes.custom_response(None, 400, "TPM-2", str(err))

        devices = BitacoraModel.get_Bitacora_by_query_csv(data)
        if not devices:
            return returnCodes.custom_response(None, 404, "TPM-4")
        
        headers = ["fechainicio", "comentario", "fechafin"]
        
        serialized_bitacora = bitacora_schema.dump(devices,many=true)
        if len(serialized_bitacora)==0:
            return returnCodes.custom_response(None, 404, "TPM-4")

        # Crear un archivo CSV en memoria
        csv_data = BytesIO()
        # Escribir las cabeceras en el CSV (codificadas como bytes)
        csv_data.write(','.join(headers).encode('utf-8') + b'\n')
        #csv_data.write(','.join(headers)+ '\n')
        for item in serialized_bitacora:
            row = [item.get("fechaInicio", ""), item.get("fechaFin", ""), item.get("comentario", "")]
            csv_data.write((('' if row[0] is None else row[0])+"," +('' if row[2] is None else row[2])+","+ ('' if row[1] is None else row[1])).encode('utf-8') + b'\n')
            #csv_data.write(','.join(str('' if value is None else value) for value in row).encode('utf-8') + b'\n')
            

    
        csv_data.seek(0)

        return send_file(
            csv_data,
            mimetype='text/csv',
            as_attachment=True,
            download_name=str(serialized_bitacora[0]['proyecto']['proyecto'])+'.csv'
        )

        #serialized_bitacora = bitacora_schema.dump(devices,many=true)
        #return returnCodes.custom_response(serialized_bitacora, 200, "TPM-3")

@nsBitacora.route("/doc")
@nsBitacora.expect(parser)
@nsBitacora.response(404, "bitacora no encontrada")
class bitacoraDOC(Resource):
    
    @nsBitacora.doc("obtener varias bitacoras")
    @api.expect(BitacoraQueryModel)
    def post(self):
        
        if request.is_json is False:
            return returnCodes.custom_response(None, 400, "TPM-2")

        req_data = request.json
        data = None
        if(not req_data):
            return returnCodes.custom_response(None, 400, "TPM-2")
        try:
            data = bitacora_schema_query.load(req_data, partial=True)
        except ValidationError as err:
            return returnCodes.custom_response(None, 400, "TPM-2", str(err))

        devices = BitacoraModel.get_Bitacora_by_query_csv(data)
        if not devices:
            return returnCodes.custom_response(None, 404, "TPM-4")
        
        serialized_bitacora = bitacora_schema.dump(devices,many=true)
        if len(serialized_bitacora)==0:
            return returnCodes.custom_response(None, 404, "TPM-4")
        
        document = Document()
        document.add_heading('Ejemplo de Documento', 0)


        table0 = document.add_table(rows=2, cols=2, style='Table Grid')
        table0.cell(0, 0).width = Inches(9.0)
        table0.cell(0, 0).text = 'fecha'
        ##table0.style = 'Colorful List'
        table0.fill_color = RGBColor(0, 0, 255)
        table0.cell(1, 0).text = 'titulo de la historia'
        table0.cell(1, 1).text = 'Libre'
        table0.cell(0, 0).merge(table0.cell(0, 1))

        paragraph = document.add_paragraph()

        paragraph.add_run('\n')

        #segunda tabla
        table1 = document.add_table(rows=2, cols=3, style='Table Grid')
        
        table1.cell(0, 0).text = 'Localizacion'
        table1.cell(0, 1).text = 'Evento'
        table1.cell(0, 2).text = 'Reality'

        table1.cell(1, 0).text = 'Alberca'
        table1.cell(1, 1).text = 'Amistad'
        table1.cell(1, 2).text = 'El bar'

        ##table0.style = 'Colorful List'
        table1.fill_color = RGBColor(0, 0, 255)
        table1.cell(1, 0).text = 'titulo de la historia'
        table1.cell(1, 1).text = 'Libre'

        paragraph = document.add_paragraph()

        paragraph.add_run('\n')

        #tercer tabla de comentarios
        table3 = document.add_table(rows=len(serialized_bitacora), cols=3, style='Table Grid')

        counter=0
        for item in serialized_bitacora:
            row = [item.get("fechaInicio", ""), item.get("fechaFin", ""), item.get("comentario", "")]
            table3.cell(counter, 0).text = '' if row[0] is None else row[0]
            table3.cell(counter, 1).text = '' if row[2] is None else row[2]
            table3.cell(counter, 2).text = '' if row[1] is None else row[1]
            counter=counter+1


        # Guardar el documento en un objeto BytesIO
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        # Devolver el archivo como respuesta con Flask
        return send_file(
            buffer,
            download_name='ejemplo34.docx',
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )